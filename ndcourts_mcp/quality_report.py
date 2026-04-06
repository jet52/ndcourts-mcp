"""Generate a Word document listing the lowest-quality opinions for Westlaw lookup.

Produces a .docx with Bluebook-formatted citations and problem descriptions,
suitable for bulk downloading replacement text from Westlaw.

Usage:
    python -m ndcourts_mcp.quality_report [--db PATH] [--count N] [--output PATH]
"""

import argparse
import re
import sqlite3
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .db import DEFAULT_DB_PATH, get_connection


def _format_bluebook_citation(cites_str: str, case_name: str, date_filed: str) -> str:
    """Format citations in Bluebook style.

    Returns: Case Name, Vol Reporter Page (Year).
    With parallel citations separated by commas.
    """
    if not cites_str:
        return case_name

    cites = [c.strip() for c in cites_str.split(";")]

    # Separate by type: ND neutral, N.D. old, N.W.2d, N.W., and skip LEXIS/WL
    nd_neutral = []
    nd_old = []
    nw2d = []
    nw = []
    for c in cites:
        if re.match(r"\d{4} ND \d+", c):
            nd_neutral.append(c)
        elif "N.D." in c and "LEXIS" not in c:
            nd_old.append(c)
        elif "N.W.2d" in c:
            nw2d.append(c)
        elif "N.W." in c and "LEXIS" not in c and "WL" not in c:
            nw.append(c)

    # Build parallel citation string: prefer ND neutral, then N.D., then N.W.2d/N.W.
    parts = []
    if nd_neutral:
        parts.append(nd_neutral[0])
    if nd_old:
        parts.append(nd_old[0])
    if nw2d:
        parts.append(nw2d[0])
    elif nw:
        parts.append(nw[0])

    if not parts:
        parts = cites[:2]  # fallback

    year = date_filed[:4] if date_filed else ""
    cite_str = ", ".join(parts)

    return f"{case_name}, {cite_str} ({year})"


def _describe_problems(row: sqlite3.Row) -> str:
    """Generate a short description of quality problems."""
    problems = []

    if row["ocr_artifacts"] > 0:
        problems.append(f"{row['ocr_artifacts']} OCR artifacts")
    if row["garbage_chars"] > 0:
        problems.append(f"{row['garbage_chars']} unrecognized characters")
    if row["has_html"]:
        problems.append("HTML/script contamination")
    if row["short_line_ratio"] > 0.15:
        problems.append(f"line-break damage ({row['short_line_ratio']:.0%} short lines)")
    elif row["short_line_ratio"] > 0.05:
        problems.append("some line-break damage")
    if row["para_markers"] == 0 and row["source_reporter"] == "ND":
        problems.append("missing paragraph markers")

    if not problems:
        problems.append("low composite score")

    return "; ".join(problems)


def generate_report(
    db_path: Path = DEFAULT_DB_PATH,
    count: int = 40,
    output: Path | None = None,
) -> Path:
    """Generate a .docx report of the lowest-quality opinions."""
    conn = get_connection(db_path)

    rows = conn.execute("""
        SELECT o.id, o.case_name, o.date_filed, o.source_reporter,
               qs.overall_score, qs.ocr_artifacts, qs.garbage_chars, qs.has_html,
               qs.short_line_ratio, qs.para_markers,
               (SELECT GROUP_CONCAT(c.citation, '; ')
                FROM citations c WHERE c.opinion_id = o.id
                ORDER BY c.is_primary DESC) as all_cites
        FROM quality_scores qs
        JOIN opinions o ON o.id = qs.opinion_id
        ORDER BY qs.overall_score ASC
        LIMIT ?
    """, (count,)).fetchall()

    conn.close()

    if output is None:
        output = Path(f"low-quality-opinions-{count}.docx")

    doc = Document()

    # Title
    title = doc.add_heading("Low-Quality Opinions — Priority Corrections", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"{count} opinions with the lowest text quality scores from the "
        f"North Dakota Supreme Court opinion database. "
        f"These opinions are candidates for replacement with clean Westlaw text."
    )
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light List Accent 1"
    table.autofit = True

    # Header row
    headers = ["#", "Citation", "Score", "Problems"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.bold = True

    # Data rows
    for idx, row in enumerate(rows, 1):
        cite_text = _format_bluebook_citation(
            row["all_cites"], row["case_name"], row["date_filed"]
        )
        problems = _describe_problems(row)
        score = f"{row['overall_score']:.0f}"

        data_row = table.add_row()
        cells = data_row.cells

        cells[0].text = str(idx)
        cells[1].text = cite_text
        cells[2].text = score
        cells[3].text = problems

        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(3.5)
        row.cells[2].width = Inches(0.5)
        row.cells[3].width = Inches(2.6)

    doc.save(str(output))
    print(f"Report saved to {output} ({count} opinions)")
    return output


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate Word doc of lowest-quality opinions for Westlaw lookup"
    )
    parser.add_argument("--db", type=Path, default=DEFAULT_DB_PATH)
    parser.add_argument("--count", type=int, default=40, help="Number of opinions (default 40)")
    parser.add_argument("--output", type=Path, default=None, help="Output .docx path")
    args = parser.parse_args()
    generate_report(args.db, args.count, args.output)


if __name__ == "__main__":
    main()
