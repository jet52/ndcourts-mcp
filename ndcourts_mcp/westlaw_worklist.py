"""Generate the daily manual-Westlaw download worklist.

The 1953-~1965 slice of the gap era (N.W.2d vols below 139) is NOT in
the court-sourced NW-cite archive, so the only path to a high-authority
text is a manual Westlaw pull. Completion of the gap era requires
acquiring every such opinion from Westlaw; that fetch is manual and
capped (~500/day). This builds the day's list as a Word doc split into
batches of <=100 citations the user can hand to Westlaw, leading with
the opinions that bear the strongest indications of error.

Target pool: gap_1953_1996 opinions still in state 'unvalidated' or
'flagged' (the court-archive promotion already drained the rest to
'corrected'/'cross_checked'), that carry an N.W. citation and are not
already on a prior worklist.

Priority (quality_score is uninformative for the NW2d OCR era per
TODO §2, so it is NOT used):
  1. flagged before unvalidated — a flag is a known court-archive
     divergence/ambiguity, the strongest error signal.
  2. missing/blank author before populated — metadata anomaly.
  3. higher inbound-citation count first — get the load-bearing
     precedent right before the obscure.
  4. older first — chip the era away chronologically.

State: every listed opinion is recorded in westlaw_requests so the
next day's run excludes it. Receiving/ingesting the returned .docs is
a separate future step (received_at/received_path columns reserved).

Usage:
  python -m ndcourts_mcp.westlaw_worklist --status
  python -m ndcourts_mcp.westlaw_worklist                 # preview 500
  python -m ndcourts_mcp.westlaw_worklist --apply          # record + .docx
  python -m ndcourts_mcp.westlaw_worklist --limit 300 --apply
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document

from .db import DEFAULT_DB_PATH, get_connection, log_provenance

DEFAULT_LIMIT = 500          # opinions per day (manual Westlaw cap)
DEFAULT_PER_SECTION = 100    # citations per Westlaw batch
WORKLIST_DIR = Path("worklists")

POOL_SQL = """
    SELECT o.id, o.case_name, o.date_filed, v.crosscheck_state,
           (o.author IS NULL OR TRIM(o.author) = '') AS missing_author,
           (SELECT c.citation FROM citations c
             WHERE c.opinion_id = o.id AND c.citation LIKE '% N.W.%'
             ORDER BY c.citation LIMIT 1) AS nw_cite,
           (SELECT COUNT(*) FROM cited_by cb
             WHERE cb.cited_opinion_id = o.id) AS inbound
    FROM validation_status v
    JOIN opinions o ON o.id = v.opinion_id
    WHERE v.era_tier = 'gap_1953_1996'
      AND v.crosscheck_state IN ('unvalidated', 'flagged')
      AND o.id NOT IN (SELECT opinion_id FROM westlaw_requests)
      AND EXISTS (SELECT 1 FROM citations c
                   WHERE c.opinion_id = o.id AND c.citation LIKE '% N.W.%')
    ORDER BY
      CASE v.crosscheck_state WHEN 'flagged' THEN 0 ELSE 1 END,
      missing_author DESC,
      inbound DESC,
      o.date_filed ASC
"""


def _ensure_table(conn) -> None:
    conn.execute("""
        CREATE TABLE IF NOT EXISTS westlaw_requests (
            opinion_id    INTEGER PRIMARY KEY REFERENCES opinions(id),
            nw_cite       TEXT NOT NULL,
            doc_batch     TEXT NOT NULL,
            listed_at     TEXT NOT NULL DEFAULT
                          (strftime('%Y-%m-%dT%H:%M:%S','now')),
            received_at   TEXT,
            received_path TEXT
        )
    """)


def _pool(conn, limit: int) -> list:
    return conn.execute(POOL_SQL + " LIMIT ?", (limit,)).fetchall()


def _write_docx(rows: list, out_path: Path, per_section: int,
                batch_date: str) -> None:
    doc = Document()
    doc.add_heading("Westlaw download worklist", level=0)
    n_sec = (len(rows) + per_section - 1) // per_section
    doc.add_paragraph(
        f"Batch date: {batch_date}.  {len(rows)} opinions in "
        f"{n_sec} section(s) of <= {per_section}.  "
        f"Look each up in Westlaw by the N.W. citation; the case name "
        f"and filing date are for your cross-reference. Ordered "
        f"highest-error-indication first."
    )
    for i in range(0, len(rows), per_section):
        chunk = rows[i:i + per_section]
        sec = i // per_section + 1
        doc.add_heading(
            f"Section {sec} of {n_sec} — {len(chunk)} citations", level=1
        )
        for r in chunk:
            flag = "  [flagged]" if r["crosscheck_state"] == "flagged" else ""
            doc.add_paragraph(
                f'{r["nw_cite"]}  —  {r["case_name"]} '
                f'({r["date_filed"]}){flag}',
                style="List Number",
            )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def status(conn) -> None:
    _ensure_table(conn)
    pool_total = conn.execute(
        "SELECT COUNT(*) FROM validation_status v JOIN opinions o "
        "ON o.id=v.opinion_id WHERE v.era_tier='gap_1953_1996' "
        "AND v.crosscheck_state IN ('unvalidated','flagged') "
        "AND EXISTS (SELECT 1 FROM citations c WHERE c.opinion_id=o.id "
        "AND c.citation LIKE '% N.W.%')"
    ).fetchone()[0]
    listed = conn.execute("SELECT COUNT(*) FROM westlaw_requests").fetchone()[0]
    received = conn.execute(
        "SELECT COUNT(*) FROM westlaw_requests WHERE received_at IS NOT NULL"
    ).fetchone()[0]
    remaining = conn.execute(
        "SELECT COUNT(*) FROM validation_status v JOIN opinions o "
        "ON o.id=v.opinion_id WHERE v.era_tier='gap_1953_1996' "
        "AND v.crosscheck_state IN ('unvalidated','flagged') "
        "AND o.id NOT IN (SELECT opinion_id FROM westlaw_requests) "
        "AND EXISTS (SELECT 1 FROM citations c WHERE c.opinion_id=o.id "
        "AND c.citation LIKE '% N.W.%')"
    ).fetchone()[0]
    print("=== Westlaw worklist status ===")
    print(f"  gap-era pool (needs Westlaw):  {pool_total}")
    print(f"  listed on prior worklists:     {listed}")
    print(f"  received back:                 {received}")
    print(f"  remaining to list:             {remaining}")


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--db", type=Path, default=DEFAULT_DB_PATH)
    ap.add_argument("--limit", type=int, default=DEFAULT_LIMIT)
    ap.add_argument("--per-section", type=int, default=DEFAULT_PER_SECTION)
    ap.add_argument("--out", type=Path)
    ap.add_argument("--apply", action="store_true",
                    help="Record opinions as listed + write the .docx "
                         "(default: preview only, records nothing)")
    ap.add_argument("--status", action="store_true")
    args = ap.parse_args()

    conn = get_connection(args.db)
    try:
        if args.status:
            status(conn)
            return
        _ensure_table(conn)
        rows = _pool(conn, args.limit)
        if not rows:
            print("Pool empty — every gap-era opinion is listed or drained.")
            return
        today = date.today().isoformat()
        n_sec = (len(rows) + args.per_section - 1) // args.per_section
        out = args.out or (WORKLIST_DIR / f"westlaw-{today}.docx")

        if not args.apply:
            print(f"DRY RUN — would list {len(rows)} opinions in {n_sec} "
                  f"section(s), write {out}. Records nothing.")
            print("  top 5:")
            for r in rows[:5]:
                print(f"    {r['nw_cite']:<18} {r['case_name'][:46]} "
                      f"[{r['crosscheck_state']}] inbound={r['inbound']}")
            print("  Pass --apply to record + generate the Word doc.")
            return

        _write_docx(rows, out, args.per_section, today)
        conn.executemany(
            "INSERT INTO westlaw_requests (opinion_id, nw_cite, doc_batch) "
            "VALUES (?, ?, ?)",
            [(r["id"], r["nw_cite"], today) for r in rows],
        )
        log_provenance(
            conn,
            operation="westlaw_worklist",
            command=f"python -m ndcourts_mcp.westlaw_worklist --apply "
                    f"--limit {args.limit}",
            rows_affected=len(rows),
            notes=(f"listed {len(rows)} gap-era opinions for manual Westlaw "
                   f"pull, batch {today}, {n_sec} sections of "
                   f"<= {args.per_section}; doc {out}"),
        )
        conn.commit()
        print(f"Listed {len(rows)} opinions (batch {today}); wrote {out}")
    finally:
        conn.close()


if __name__ == "__main__":
    main()
